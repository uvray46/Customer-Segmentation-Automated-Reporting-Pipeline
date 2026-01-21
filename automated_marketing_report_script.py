
import pandas as pd
import matplotlib.pyplot as plt
from pathlib import Path
from docx import Document
from docx.shared import Inches

# Step 1: Load your dataset
df = pd.read_csv("bank.csv")

# Step 2: Add calculated columns
df['conversion'] = df['deposit'].map({'yes': 1, 'no': 0})
df['has_previous_contact'] = df['pdays'].apply(lambda x: 1 if x != -1 else 0)

# Step 3: Group and summarize
job_summary = df.groupby('job').agg({
    'age': 'mean',
    'balance': 'mean',
    'campaign': 'mean',
    'conversion': 'mean',
    'has_previous_contact': 'mean'
}).rename(columns={
    'age': 'avg_age',
    'balance': 'avg_balance',
    'campaign': 'avg_campaign_contacts',
    'conversion': 'conversion_rate',
    'has_previous_contact': 'prior_contact_rate'
}).reset_index()

# Step 4: Targeting Logic
avg_conversion = job_summary['conversion_rate'].mean()
avg_balance = job_summary['avg_balance'].mean()

def targeting_strategy(row):
    if row['conversion_rate'] > avg_conversion and row['avg_balance'] > avg_balance:
        return "High Priority Target"
    elif row['conversion_rate'] > avg_conversion:
        return "Consider – Good Conversion"
    elif row['avg_balance'] > avg_balance:
        return "Consider – Strong Financial Profile"
    else:
        return "Low Priority"

job_summary['targeting_recommendation'] = job_summary.apply(targeting_strategy, axis=1)

# Step 5: Save to Excel
excel_path = Path("automated_marketing_report.xlsx")
with pd.ExcelWriter(excel_path, engine='xlsxwriter') as writer:
    job_summary.to_excel(writer, sheet_name='Job Summary', index=False)
    workbook = writer.book
    worksheet = writer.sheets['Job Summary']
    worksheet.set_column('A:H', 30)
    worksheet.write('J1', "Targeting Recommendation Key:")
    worksheet.write('J2', "High Priority Target – high conversion + high balance")
    worksheet.write('J3', "Consider – Good Conversion or Financial Profile")
    worksheet.write('J4', "Low Priority – lower conversion and financial potential")

    chart = workbook.add_chart({'type': 'column'})
    chart.add_series({
        'name': 'Conversion Rate',
        'categories': ['Job Summary', 1, 0, len(job_summary), 0],
        'values':     ['Job Summary', 1, 1, len(job_summary), 1],
    })
    chart.add_series({
        'name': 'Average Balance',
        'categories': ['Job Summary', 1, 0, len(job_summary), 0],
        'values':     ['Job Summary', 1, 2, len(job_summary), 2],
        'y2_axis':    True,
    })
    chart.set_title({'name': 'Conversion Rate & Average Balance by Job'})
    chart.set_x_axis({'name': 'Job'})
    chart.set_y_axis({'name': 'Conversion Rate'})
    chart.set_y2_axis({'name': 'Average Balance'})
    worksheet.insert_chart('B15', chart)

# Step 6: Create Chart Image for Word Report
chart_path = Path("job_performance_chart.png")
plt.figure(figsize=(10, 6))
ax1 = plt.gca()
job_summary.plot(kind='bar', x='job', y='conversion_rate', ax=ax1, color='skyblue', legend=False)
ax2 = ax1.twinx()
job_summary.plot(kind='line', x='job', y='avg_balance', ax=ax2, color='darkred', marker='o', legend=False)

ax1.set_ylabel("Conversion Rate")
ax2.set_ylabel("Avg Balance")
ax1.set_title("Conversion Rate vs Avg Balance by Job")
ax1.set_xticklabels(job_summary['job'], rotation=45, ha='right')

plt.tight_layout()
plt.savefig(chart_path)
plt.close()

# Step 7: Create Word Report
doc = Document()
doc.add_heading('Marketing Insights Report', 0)

doc.add_paragraph(
    "This report summarizes key insights from the marketing dataset analysis. "
    "The goal was to identify the most valuable customer segments and recommend targeting strategies based on conversion rates and financial profiles."
)

doc.add_heading('Summary of Findings', level=1)
doc.add_paragraph(
    f"- Average Conversion Rate Across All Jobs: {avg_conversion:.2%}\n"
    f"- Average Balance Across All Jobs: ${avg_balance:,.2f}\n"
)
doc.add_paragraph("Roles with the highest marketing value based on above-average conversion rates and balances were classified as 'High Priority Targets'. Roles with one strong metric were marked for consideration.")

doc.add_heading('Targeting Recommendations by Job Role', level=1)
for _, row in job_summary.iterrows():
    doc.add_paragraph(
        f"{row['job'].capitalize()}: {row['targeting_recommendation']}\n"
        f"  - Conversion Rate: {row['conversion_rate']:.2%}\n"
        f"  - Avg. Balance: ${row['avg_balance']:,.2f}\n"
        f"  - Campaign Contacts: {row['avg_campaign_contacts']:.1f}\n",
        style='List Bullet'
    )

doc.add_heading("Conversion Rate & Average Balance by Job", level=2)
doc.add_picture(str(chart_path), width=Inches(6))

doc.add_paragraph("The chart above compares the conversion rate and average balance for each job segment. "
                  "Higher conversion segments such as Management and Admin roles may represent strong targets for future campaigns.")

doc_path = Path("automated_marketing_report.docx")
doc.save(doc_path)

print("Reports generated successfully:")
print(f"Excel: {excel_path}")
print(f"Word: {doc_path}")
